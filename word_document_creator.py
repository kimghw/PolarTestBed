"""
MS Word 문서 생성 및 마크다운 스타일 자동 적용 스크립트

이 스크립트는 마크다운 형식의 텍스트를 읽어서 Word 문서로 변환하며,
다음과 같은 마크다운 요소를 Word 스타일로 자동 변환합니다:
- # (제목1)
- ## (제목2)
- ### (제목3)
- 1. (번호 목록)
- - (불릿 목록)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re


class MarkdownToWordConverter:
    """마크다운 텍스트를 Word 문서로 변환하는 클래스"""

    def __init__(self):
        """문서 초기화 및 스타일 설정"""
        self.doc = Document()
        self._setup_styles()
        self.current_list_level = 0

    def _setup_styles(self):
        """Word 문서에 사용할 스타일 정의"""
        styles = self.doc.styles

        # 제목1 스타일 (# 에 대응)
        try:
            heading1 = styles['Heading 1']
        except KeyError:
            heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1.font.name = '맑은 고딕'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 139)  # 진한 파란색
        heading1.paragraph_format.space_before = Pt(12)
        heading1.paragraph_format.space_after = Pt(6)

        # 제목2 스타일 (## 에 대응)
        try:
            heading2 = styles['Heading 2']
        except KeyError:
            heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2.font.name = '맑은 고딕'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 0, 100)  # 파란색
        heading2.paragraph_format.space_before = Pt(12)
        heading2.paragraph_format.space_after = Pt(6)

        # 제목3 스타일 (### 에 대응)
        try:
            heading3 = styles['Heading 3']
        except KeyError:
            heading3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        heading3.font.name = '맑은 고딕'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.font.color.rgb = RGBColor(50, 50, 50)  # 진한 회색
        heading3.paragraph_format.space_before = Pt(6)
        heading3.paragraph_format.space_after = Pt(3)

    def _detect_line_type(self, line):
        """라인의 마크다운 타입을 감지"""
        line = line.strip()

        if not line:
            return 'empty', line

        # 구분선 감지
        if line == '---':
            return 'separator', line

        # 제목 레벨 감지
        if line.startswith('###'):
            return 'heading3', line[3:].strip()
        elif line.startswith('##'):
            return 'heading2', line[2:].strip()
        elif line.startswith('#'):
            return 'heading1', line[1:].strip()

        # 번호 목록 감지
        numbered_pattern = r'^(\d+)\.\s+(.+)$'
        match = re.match(numbered_pattern, line)
        if match:
            return 'numbered', match.group(2)

        # 불릿 목록 감지
        if line.startswith('- '):
            return 'bullet', line[2:].strip()

        # 들여쓰기된 불릿 목록 감지
        indented_bullet = r'^(\s+)-\s+(.+)$'
        match = re.match(indented_bullet, line)
        if match:
            indent_level = len(match.group(1)) // 2  # 2칸을 1레벨로
            return 'bullet_indent', (match.group(2), indent_level)

        return 'normal', line

    def add_line_to_document(self, line):
        """라인을 분석하여 적절한 스타일로 문서에 추가"""
        line_type, content = self._detect_line_type(line)

        if line_type == 'empty':
            # 빈 줄은 무시 (필요시 단락 구분용으로 사용 가능)
            pass

        elif line_type == 'separator':
            # 구분선 추가
            p = self.doc.add_paragraph('_' * 50)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.style.font.color.rgb = RGBColor(192, 192, 192)

        elif line_type == 'heading1':
            p = self.doc.add_paragraph(content, style='Heading 1')

        elif line_type == 'heading2':
            p = self.doc.add_paragraph(content, style='Heading 2')

        elif line_type == 'heading3':
            p = self.doc.add_paragraph(content, style='Heading 3')

        elif line_type == 'numbered':
            # 번호 목록
            p = self.doc.add_paragraph(style='List Number')
            p.add_run(content)
            p.style.font.name = '맑은 고딕'
            p.style.font.size = Pt(11)

        elif line_type == 'bullet':
            # 불릿 목록
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(content)
            p.style.font.name = '맑은 고딕'
            p.style.font.size = Pt(11)

        elif line_type == 'bullet_indent':
            # 들여쓰기된 불릿 목록
            content_text, indent_level = content
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(content_text)
            p.style.font.name = '맑은 고딕'
            p.style.font.size = Pt(11)
            # 들여쓰기 레벨 적용
            p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))

        else:  # normal
            # 일반 텍스트
            p = self.doc.add_paragraph(content)
            p.style.font.name = '맑은 고딕'
            p.style.font.size = Pt(11)

    def convert_markdown_file(self, markdown_file_path):
        """마크다운 파일을 읽어서 Word 문서로 변환"""
        try:
            with open(markdown_file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()

            for line in lines:
                self.add_line_to_document(line)

            print(f"✅ 성공적으로 {len(lines)}개의 라인을 처리했습니다.")

        except FileNotFoundError:
            print(f"❌ 파일을 찾을 수 없습니다: {markdown_file_path}")
        except Exception as e:
            print(f"❌ 파일 처리 중 오류 발생: {e}")

    def save_document(self, output_path):
        """Word 문서를 파일로 저장"""
        try:
            self.doc.save(output_path)
            print(f"✅ Word 문서가 '{output_path}'로 저장되었습니다.")
        except Exception as e:
            print(f"❌ 문서 저장 중 오류 발생: {e}")


def main():
    """메인 실행 함수"""
    # 변환기 인스턴스 생성
    converter = MarkdownToWordConverter()

    # 마크다운 파일 경로
    markdown_file = '목차.md'

    # Word 문서 출력 경로
    output_file = '극지_친환경_추진시스템_시험베드_목차.docx'

    print("=" * 60)
    print("📄 마크다운 → Word 문서 변환 시작")
    print("=" * 60)

    # 마크다운 파일을 변환
    converter.convert_markdown_file(markdown_file)

    # Word 문서 저장
    converter.save_document(output_file)

    print("=" * 60)
    print("✨ 변환 완료!")
    print("=" * 60)

    # 스타일 적용 정보 출력
    print("\n📋 적용된 스타일 정보:")
    print("  - # → 제목1 (16pt, 진한 파란색, 굵게)")
    print("  - ## → 제목2 (14pt, 파란색, 굵게)")
    print("  - ### → 제목3 (12pt, 진한 회색, 굵게)")
    print("  - 1. → 번호 목록")
    print("  - - → 불릿 목록")
    print("  - --- → 구분선")
    print("  - 들여쓰기된 불릿도 자동 감지 및 적용")


if __name__ == "__main__":
    main()